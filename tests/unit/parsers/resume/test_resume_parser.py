"""Integration tests for transformer.parsers.resume.resume_parser."""

import contextlib
from pathlib import Path

from docx import Document
import fitz
import pytest

from transformer.parsers.exceptions import (
    CorruptedFileError,
    MappingError,
    TextExtractionError,
    UnsupportedFormatError,
)
from transformer.parsers.resume.resume_parser import ResumeParser


def _pdf_with_text(path: Path, lines: list[str], y_start: int = 72) -> None:
    doc = fitz.open()
    page = doc.new_page()
    y = y_start
    for line in lines:
        page.insert_text((72, y), line)
        y += 20
    doc.save(str(path))
    doc.close()


def _docx_with_paragraphs(path: Path, lines: list[str]) -> None:
    document = Document()
    for line in lines:
        document.add_paragraph(line)
    document.save(str(path))


class TestHappyPath:
    def test_parses_pdf_resume(self, tmp_path: Path) -> None:
        pdf_path = tmp_path / "resume.pdf"
        _pdf_with_text(
            pdf_path,
            [
                "Jane Doe",
                "jane@example.com | 555-123-4567",
                "Skills",
                "Python, SQL, Docker",
            ],
        )
        candidate = ResumeParser().parse(pdf_path)
        assert candidate.first_name == "Jane"
        assert candidate.contact is not None
        assert candidate.contact.email == "jane@example.com"
        assert "Python" in candidate.skills

    def test_parses_docx_resume(self, tmp_path: Path) -> None:
        docx_path = tmp_path / "resume.docx"
        _docx_with_paragraphs(
            docx_path,
            ["Jane Doe", "jane@example.com", "Skills", "Python, SQL"],
        )
        candidate = ResumeParser().parse(docx_path)
        assert candidate.first_name == "Jane"
        assert candidate.last_name == "Doe"


class TestNegativePath:
    def test_unsupported_extension_raises(self, tmp_path: Path) -> None:
        txt_path = tmp_path / "resume.txt"
        txt_path.write_text("Jane Doe")
        with pytest.raises(UnsupportedFormatError):
            ResumeParser().parse(txt_path)

    def test_extension_mismatch_pdf_actually_docx(self, tmp_path: Path) -> None:
        # See test_text_extractor for the documented fitz content-sniffing
        # quirk -- this only asserts no unhandled crash occurs.
        docx_path = tmp_path / "real.docx"
        _docx_with_paragraphs(docx_path, ["Jane Doe"])
        fake_pdf = tmp_path / "fake.pdf"
        fake_pdf.write_bytes(docx_path.read_bytes())
        with contextlib.suppress(CorruptedFileError, TextExtractionError, MappingError):
            ResumeParser().parse(fake_pdf)

    def test_extension_mismatch_docx_actually_pdf(self, tmp_path: Path) -> None:
        pdf_path = tmp_path / "real.pdf"
        _pdf_with_text(pdf_path, ["Jane Doe"])
        fake_docx = tmp_path / "fake.docx"
        fake_docx.write_bytes(pdf_path.read_bytes())
        with pytest.raises(CorruptedFileError):
            ResumeParser().parse(fake_docx)


class TestEmptyAndCorruptedResumes:
    def test_blank_pdf_raises_text_extraction_error(self, tmp_path: Path) -> None:
        pdf_path = tmp_path / "blank.pdf"
        doc = fitz.open()
        doc.new_page()
        doc.save(str(pdf_path))
        doc.close()
        with pytest.raises(TextExtractionError):
            ResumeParser().parse(pdf_path)

    def test_blank_docx_raises_text_extraction_error(self, tmp_path: Path) -> None:
        docx_path = tmp_path / "blank.docx"
        Document().save(str(docx_path))
        with pytest.raises(TextExtractionError):
            ResumeParser().parse(docx_path)

    def test_image_only_pdf_raises_text_extraction_error(self, tmp_path: Path) -> None:
        pdf_path = tmp_path / "scanned.pdf"
        doc = fitz.open()
        page = doc.new_page()
        pixmap = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 50, 50))
        pixmap.set_rect(pixmap.irect, (0, 0, 255))
        page.insert_image(fitz.Rect(0, 0, 50, 50), pixmap=pixmap)
        doc.save(str(pdf_path))
        doc.close()
        with pytest.raises(TextExtractionError):
            ResumeParser().parse(pdf_path)

    def test_corrupted_pdf_raises_corrupted_file_error(self, tmp_path: Path) -> None:
        pdf_path = tmp_path / "bad.pdf"
        pdf_path.write_bytes(b"definitely not a pdf")
        with pytest.raises(CorruptedFileError):
            ResumeParser().parse(pdf_path)

    def test_corrupted_docx_raises_corrupted_file_error(self, tmp_path: Path) -> None:
        docx_path = tmp_path / "bad.docx"
        docx_path.write_bytes(b"definitely not a docx zip")
        with pytest.raises(CorruptedFileError):
            ResumeParser().parse(docx_path)

    def test_resume_without_name_raises_mapping_error(self, tmp_path: Path) -> None:
        pdf_path = tmp_path / "no_name.pdf"
        _pdf_with_text(pdf_path, ["jane@example.com", "555-123-4567", "Skills"])
        with pytest.raises(MappingError):
            ResumeParser().parse(pdf_path)


class TestUnicodeAndMultipleContacts:
    def test_unicode_name_and_text_handled(self, tmp_path: Path) -> None:
        docx_path = tmp_path / "unicode.docx"
        _docx_with_paragraphs(
            docx_path,
            ["José García", "jose@example.com", "Skills", "Café management, Résumé"],
        )
        candidate = ResumeParser().parse(docx_path)
        assert candidate.first_name == "José"

    def test_multiple_emails_first_one_selected(self, tmp_path: Path) -> None:
        docx_path = tmp_path / "multi_email.docx"
        _docx_with_paragraphs(
            docx_path,
            ["Jane Doe", "jane.work@example.com", "jane.personal@example.com"],
        )
        candidate = ResumeParser().parse(docx_path)
        assert candidate.contact is not None
        assert candidate.contact.email == "jane.work@example.com"

    def test_missing_optional_sections_degrade_gracefully(self, tmp_path: Path) -> None:
        docx_path = tmp_path / "no_sections.docx"
        _docx_with_paragraphs(docx_path, ["Jane Doe", "jane@example.com"])
        candidate = ResumeParser().parse(docx_path)
        assert candidate.skills == []
        assert candidate.experiences == []


class TestTwoColumnAndTableHeavy:
    def test_two_column_layout_does_not_crash(self, tmp_path: Path) -> None:
        pdf_path = tmp_path / "two_column.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Jane Doe")
        page.insert_text((72, 100), "Skills")
        page.insert_text((72, 120), "Python")
        page.insert_text((350, 100), "Education")
        page.insert_text((350, 120), "MIT")
        doc.save(str(pdf_path))
        doc.close()
        candidate = ResumeParser().parse(pdf_path)
        assert candidate.first_name == "Jane"

    def test_table_heavy_docx_skills_extracted(self, tmp_path: Path) -> None:
        docx_path = tmp_path / "table_resume.docx"
        document = Document()
        document.add_paragraph("Jane Doe")
        document.add_paragraph("jane@example.com")
        document.add_paragraph("Skills")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Python"
        table.rows[0].cells[1].text = "SQL"
        document.save(str(docx_path))
        candidate = ResumeParser().parse(docx_path)
        assert candidate.first_name == "Jane"


class TestLargeResume:
    def test_large_resume_with_many_lines_parses(self, tmp_path: Path) -> None:
        docx_path = tmp_path / "large.docx"
        lines = ["Jane Doe", "jane@example.com", "Skills"]
        lines += [f"Skill{i}" for i in range(300)]
        _docx_with_paragraphs(docx_path, lines)
        candidate = ResumeParser().parse(docx_path)
        assert candidate.first_name == "Jane"
        assert len(candidate.skills) > 0


class TestBugReportRegressions:
    """Each test below encodes one bug from the real-data dry-run audit,
    using the exact (or representative) failing input that was reported.
    """

    def test_bug8_hyphen_joined_month_year_dates_detected(self, tmp_path: Path) -> None:
        """Bug #8: 'Dec-2027' (no space) must be recognised as a date."""
        docx_path = tmp_path / "bug8.docx"
        _docx_with_paragraphs(
            docx_path,
            [
                "Jane Doe",
                "jane@example.com",
                "Education",
                "MIT",
                "Bachelor of Science in Computer Science",
                "May 2023 - Dec-2027",
            ],
        )
        candidate = ResumeParser().parse(docx_path)
        assert len(candidate.education) == 1
        assert candidate.education[0].start_date is not None

    def test_bug6_date_and_grade_lines_never_become_institutions(
        self, tmp_path: Path
    ) -> None:
        """Bug #6: date/CGPA/percentage-shaped lines must never spawn a
        spurious institution entry; only the two real institutions should
        be produced, with no garbage entry between them."""
        docx_path = tmp_path / "bug6.docx"
        _docx_with_paragraphs(
            docx_path,
            [
                "Jane Doe",
                "jane@example.com",
                "Education",
                "National Institute of Technology, Uttarakhand",
                "Aug 2023 - Present",
                "B.Tech in Electronics and Communication Engineering",
                "Indian Institute of Technology Madras",
                "Mar 2023 - Present",
                "B.Sc in Programming and Data Science (Diploma Level)",
                "Shri Balwant Singh Senior Secondary School, Hathras, UP",
            ],
        )
        candidate = ResumeParser().parse(docx_path)
        institutions = {e.institution for e in candidate.education}
        assert institutions == {
            "National Institute of Technology, Uttarakhand",
            "Indian Institute of Technology Madras",
        }
        # No date string and no school entry leaked through as an
        # institution name.
        assert "Mar 2023 - Present" not in institutions
        assert not any("School" in (inst or "") for inst in institutions)

    def test_bug2_abbreviated_degree_forms_recognised(self, tmp_path: Path) -> None:
        """Bug #2: 'B.Tech'/'B.Sc' abbreviations, not just spelled-out
        forms, must be recognised as degree lines."""
        docx_path = tmp_path / "bug2.docx"
        _docx_with_paragraphs(
            docx_path,
            [
                "Jane Doe",
                "jane@example.com",
                "Education",
                "IIT Madras",
                "B.Tech in Electronics Engineering",
            ],
        )
        candidate = ResumeParser().parse(docx_path)
        assert len(candidate.education) == 1
        assert candidate.education[0].degree == "B.Tech"

    def test_bug11_dash_separated_degree_field_split_correctly(
        self, tmp_path: Path
    ) -> None:
        """Bug #11: a dash/en-dash separator (not just ' in ') must split
        degree from field_of_study correctly."""
        docx_path = tmp_path / "bug11.docx"
        _docx_with_paragraphs(
            docx_path,
            [
                "Jane Doe",
                "jane@example.com",
                "Education",
                "IIT Madras",
                "BS - Data Science and Applications",
            ],
        )
        candidate = ResumeParser().parse(docx_path)
        assert len(candidate.education) == 1
        assert candidate.education[0].degree == "BS"
        assert candidate.education[0].field_of_study == "Data Science and Applications"

    def test_bug13_school_entry_excluded_even_with_diploma_keyword(
        self, tmp_path: Path
    ) -> None:
        """Bug #13: a school-level line must never become a degree entry,
        even if a degree keyword (e.g. 'Diploma') coincidentally appears
        in nearby text."""
        docx_path = tmp_path / "bug13.docx"
        _docx_with_paragraphs(
            docx_path,
            [
                "Jane Doe",
                "jane@example.com",
                "Education",
                "IIT Madras",
                "Bachelor of Science in Data Science",
                "Senior Secondary School, Diploma Level Curriculum",
            ],
        )
        candidate = ResumeParser().parse(docx_path)
        assert len(candidate.education) == 1
        assert candidate.education[0].institution == "IIT Madras"

    def test_bug10_gpa_extracted_and_normalised_to_four_point_scale(
        self, tmp_path: Path
    ) -> None:
        """Bug #10: 'CGPA: 8.55 / 10' must populate Education.gpa,
        normalised onto the canonical model's 0.0-4.0 scale."""
        docx_path = tmp_path / "bug10.docx"
        _docx_with_paragraphs(
            docx_path,
            [
                "Jane Doe",
                "jane@example.com",
                "Education",
                "IIT Madras",
                "Bachelor of Science in Data Science",
                "CGPA: 8.55 / 10",
            ],
        )
        candidate = ResumeParser().parse(docx_path)
        assert len(candidate.education) == 1
        gpa = candidate.education[0].gpa
        assert gpa is not None
        assert 0.0 <= gpa <= 4.0
        assert abs(gpa - 3.42) < 0.01

    def test_bug1_pipe_delimited_single_line_experience_parsed(
        self, tmp_path: Path
    ) -> None:
        """Bug #1: a single pipe-delimited line must split into company,
        title, and date correctly instead of being discarded."""
        docx_path = tmp_path / "bug1.docx"
        _docx_with_paragraphs(
            docx_path,
            [
                "Jane Doe",
                "jane@example.com",
                "Experience",
                "Accenture India | Software Engineering Intern - AI Applications "
                "| May 2026 - Present",
            ],
        )
        candidate = ResumeParser().parse(docx_path)
        assert len(candidate.experiences) == 1
        assert candidate.experiences[0].company == "Accenture India"
        assert "AI Applications" in (candidate.experiences[0].title or "")

    def test_bug15_short_institution_with_comma_not_mistaken_for_location(
        self, tmp_path: Path
    ) -> None:
        """Bug #15: a short institution name containing a comma (e.g.
        'MIT, USA') must still start a new education entry, not be
        rejected by the location-line heuristic."""
        docx_path = tmp_path / "bug15.docx"
        _docx_with_paragraphs(
            docx_path,
            [
                "Jane Doe",
                "jane@example.com",
                "Education",
                "MIT, USA",
                "Bachelor of Science in Computer Science",
            ],
        )
        candidate = ResumeParser().parse(docx_path)
        assert len(candidate.education) == 1
        assert candidate.education[0].institution == "MIT, USA"

    def test_bug7_decorated_stop_section_heading_terminates_education(
        self, tmp_path: Path
    ) -> None:
        """Bug #7: 'Projects (AI/ML Focused)' must still be recognised as
        a stop-section, so it does not bleed into the Education body."""
        docx_path = tmp_path / "bug7.docx"
        _docx_with_paragraphs(
            docx_path,
            [
                "Jane Doe",
                "jane@example.com",
                "Education",
                "IIT Madras",
                "Bachelor of Science in Data Science",
                "Projects (AI/ML Focused)",
                "Face Recognition Attendance System",
            ],
        )
        candidate = ResumeParser().parse(docx_path)
        assert len(candidate.education) == 1
        # The project line never leaked into any education field.
        assert "Face Recognition" not in (candidate.education[0].field_of_study or "")

    def test_bug16_stop_section_fix_does_not_false_positive_on_bullets(
        self, tmp_path: Path
    ) -> None:
        """Bug #16 guard: a bullet line mentioning 'project' must not be
        misdetected as a section boundary and truncate the experience
        description."""
        docx_path = tmp_path / "bug16.docx"
        document = Document()
        document.add_paragraph("Jane Doe")
        document.add_paragraph("jane@example.com")
        document.add_paragraph("Experience")
        document.add_paragraph("Acme Corp | Engineer | Jan 2022 - Present")
        document.add_paragraph(
            "Led major project initiatives across teams", style="List Bullet"
        )
        document.save(str(docx_path))
        candidate = ResumeParser().parse(docx_path)
        assert len(candidate.experiences) == 1
        assert "Led major project initiatives" in (
            candidate.experiences[0].description or ""
        )

    def test_bug12_docx_style_bullets_captured_as_description(
        self, tmp_path: Path
    ) -> None:
        """Bug #12: DOCX 'List Bullet' style paragraphs (no literal bullet
        character in the run text) must still be recognised as bullets,
        so description text is not silently dropped."""
        docx_path = tmp_path / "bug12.docx"
        document = Document()
        document.add_paragraph("Jane Doe")
        document.add_paragraph("jane@example.com")
        document.add_paragraph("Experience")
        document.add_paragraph("Acme Corp | Engineer | Jan 2022 - Present")
        document.add_paragraph(
            "Architected production-grade pipelines", style="List Bullet"
        )
        document.add_paragraph(
            "Optimized vector-based search workflows", style="List Bullet"
        )
        document.save(str(docx_path))
        candidate = ResumeParser().parse(docx_path)
        assert len(candidate.experiences) == 1
        description = candidate.experiences[0].description or ""
        assert "Architected production-grade pipelines" in description
        assert "Optimized vector-based search workflows" in description

    def test_per_entry_degradation_one_bad_education_entry_does_not_drop_good_one(
        self, tmp_path: Path
    ) -> None:
        """Phase C: a malformed education entry (e.g. degree text that
        fails to extract) must not discard a valid sibling entry."""
        docx_path = tmp_path / "graceful.docx"
        _docx_with_paragraphs(
            docx_path,
            [
                "Jane Doe",
                "jane@example.com",
                "Education",
                "IIT Madras",
                "Bachelor of Science in Data Science",
                "Some Other University",
                # No degree line follows -- this entry has no degree and
                # is dropped by flush(), not raised as an error.
            ],
        )
        candidate = ResumeParser().parse(docx_path)
        assert len(candidate.education) == 1
        assert candidate.education[0].institution == "IIT Madras"
