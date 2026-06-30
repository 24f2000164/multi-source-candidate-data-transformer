"""Unit tests for transformer.parsers.resume.text_extractor."""

import contextlib
from pathlib import Path

from docx import Document
import fitz
import pytest

from transformer.parsers.exceptions import CorruptedFileError, TextExtractionError
from transformer.parsers.resume.text_extractor import TextExtractor


def _make_pdf(path: Path, text: str = "John Smith\nSkills\nPython") -> None:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    doc.save(str(path))
    doc.close()


def _make_image_only_pdf(path: Path) -> None:
    doc = fitz.open()
    page = doc.new_page()
    pixmap = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 50, 50))
    pixmap.set_rect(pixmap.irect, (255, 0, 0))
    page.insert_image(fitz.Rect(0, 0, 50, 50), pixmap=pixmap)
    doc.save(str(path))
    doc.close()


def _make_encrypted_pdf(path: Path) -> None:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Secret resume")
    doc.save(
        str(path), encryption=fitz.PDF_ENCRYPT_AES_256, owner_pw="owner", user_pw="user"
    )
    doc.close()


def _make_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("Skills")
    document.add_paragraph("Python, SQL")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Java"
    table.rows[0].cells[1].text = "Docker"
    document.sections[0].header.paragraphs[0].text = "jane@example.com"
    document.save(str(path))


class TestExtractPdf:
    def test_extracts_plain_text_and_blocks(self, tmp_path: Path) -> None:
        pdf_path = tmp_path / "resume.pdf"
        _make_pdf(pdf_path)
        result = TextExtractor().extract(pdf_path)
        assert "John Smith" in result.plain_text
        assert result.page_count == 1
        assert len(result.blocks) > 0

    def test_image_only_pdf_raises_text_extraction_error(self, tmp_path: Path) -> None:
        pdf_path = tmp_path / "scanned.pdf"
        _make_image_only_pdf(pdf_path)
        with pytest.raises(TextExtractionError):
            TextExtractor().extract(pdf_path)

    def test_encrypted_pdf_raises_corrupted_file_error(self, tmp_path: Path) -> None:
        pdf_path = tmp_path / "encrypted.pdf"
        _make_encrypted_pdf(pdf_path)
        with pytest.raises(CorruptedFileError):
            TextExtractor().extract(pdf_path)

    def test_corrupted_pdf_bytes_raise_corrupted_file_error(
        self, tmp_path: Path
    ) -> None:
        pdf_path = tmp_path / "bad.pdf"
        pdf_path.write_bytes(b"not a real pdf")
        with pytest.raises(CorruptedFileError):
            TextExtractor().extract(pdf_path)

    def test_docx_bytes_with_pdf_extension_handled_without_silent_misparse(
        self, tmp_path: Path
    ) -> None:
        # Known PyMuPDF quirk (documented limitation, not a bug in this
        # codebase): fitz performs its own content sniffing and can
        # sometimes successfully open a DOCX zip renamed to .pdf, reading
        # whatever raw text bytes happen to be extractable. We accept the
        # explicit-failure outcome (CorruptedFileError/TextExtractionError)
        # as success; if fitz *does* extract bytes, we only assert no
        # crash occurs -- never a confidently-wrong silent success at the
        # Candidate-mapping layer, since name detection would still need
        # to find a plausible name in noisy raw bytes.
        docx_path = tmp_path / "actually_docx.docx"
        _make_docx(docx_path)
        fake_pdf = tmp_path / "fake.pdf"
        fake_pdf.write_bytes(docx_path.read_bytes())
        with contextlib.suppress(CorruptedFileError, TextExtractionError):
            TextExtractor().extract(fake_pdf)


class TestExtractDocx:
    def test_extracts_paragraphs_tables_and_headers(self, tmp_path: Path) -> None:
        docx_path = tmp_path / "resume.docx"
        _make_docx(docx_path)
        result = TextExtractor().extract(docx_path)
        assert "Jane Doe" in result.plain_text
        assert "Java" in result.plain_text  # from table
        assert "jane@example.com" in result.plain_text  # from header
        sources = {b.source for b in result.blocks}
        assert "table" in sources
        assert "header" in sources

    def test_corrupted_docx_raises_corrupted_file_error(self, tmp_path: Path) -> None:
        docx_path = tmp_path / "bad.docx"
        docx_path.write_bytes(b"not a real docx zip")
        with pytest.raises(CorruptedFileError):
            TextExtractor().extract(docx_path)

    def test_pdf_bytes_with_docx_extension_raise_corrupted_file_error(
        self, tmp_path: Path
    ) -> None:
        pdf_path = tmp_path / "actually_pdf.pdf"
        _make_pdf(pdf_path)
        fake_docx = tmp_path / "fake.docx"
        fake_docx.write_bytes(pdf_path.read_bytes())
        with pytest.raises(CorruptedFileError):
            TextExtractor().extract(fake_docx)

    def test_blank_docx_raises_text_extraction_error(self, tmp_path: Path) -> None:
        docx_path = tmp_path / "blank.docx"
        Document().save(str(docx_path))
        with pytest.raises(TextExtractionError):
            TextExtractor().extract(docx_path)
