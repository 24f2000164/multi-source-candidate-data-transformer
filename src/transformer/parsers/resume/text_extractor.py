"""PDF/DOCX -> raw text + block metadata.

Responsibility: ONLY get text (plain + block-structured where available) out
of a binary file. No section/field parsing logic lives here.
"""

from dataclasses import dataclass, field
import logging
from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import fitz  # PyMuPDF

from transformer.parsers.exceptions import CorruptedFileError, TextExtractionError

logger = logging.getLogger(__name__)

_MIN_USABLE_TEXT_CHARS = 10


@dataclass(frozen=True)
class TextBlock:
    """A position-aware (PDF) or paragraph-level (DOCX) span of text.

    Attributes:
        text: The block's raw text content.
        source: Origin of the block (``"body"``, ``"table"``, ``"header"``,
            ``"footer"``).
        page_index: Zero-based page number (PDF only; ``0`` for DOCX).
        order: Reading-order position within the document, for "first
            block" heuristics.
        is_isolated: ``True`` when the block sits alone (its own visual
            unit) rather than embedded inline within a larger paragraph --
            a stronger header/name signal when available.
    """

    text: str
    source: str = "body"
    page_index: int = 0
    order: int = 0
    is_isolated: bool = True


@dataclass(frozen=True)
class ExtractedText:
    """Container for everything extracted from a resume source file.

    Attributes:
        plain_text: Full document text, newline-joined.
        blocks: Position/structure-aware spans (see ``TextBlock``).
        page_count: Number of pages (PDF) or ``1`` for DOCX.
    """

    plain_text: str
    blocks: list[TextBlock] = field(default_factory=list)
    page_count: int = 1


class TextExtractor:
    """Extracts plain text and structural blocks from PDF/DOCX resumes."""

    def extract(self, source: Path) -> ExtractedText:
        """Extract text from ``source`` based on its file extension.

        Args:
            source: Path to a ``.pdf`` or ``.docx`` resume file.

        Returns:
            The extracted text and block metadata.

        Raises:
            CorruptedFileError: If the file cannot be opened by its
                declared format (corrupted, encrypted, or an
                extension/content mismatch).
            TextExtractionError: If no usable text could be extracted
                (e.g. an image-only PDF).
        """
        suffix = source.suffix.lower()
        if suffix == ".pdf":
            extracted = self._extract_pdf(source)
        elif suffix == ".docx":
            extracted = self._extract_docx(source)
        else:
            # FileValidator should already reject this; defensive guard.
            raise CorruptedFileError(f"Unsupported file extension: {suffix}")

        if len(extracted.plain_text.strip()) < _MIN_USABLE_TEXT_CHARS:
            logger.warning("resume_no_usable_text", extra={"path": str(source)})
            raise TextExtractionError(
                f"No extractable text found in resume: {source.name}"
            )
        return extracted

    def _extract_pdf(self, source: Path) -> ExtractedText:
        """Extract text and blocks from a PDF, single-pass per page.

        Args:
            source: Path to the PDF file.

        Returns:
            Extracted text with PDF block metadata.

        Raises:
            CorruptedFileError: If the PDF is corrupted, encrypted, or not
                actually a PDF (extension/content mismatch).
        """
        try:
            doc = fitz.open(source)
        except Exception as exc:
            logger.warning("resume_pdf_open_failed", extra={"path": str(source)})
            raise CorruptedFileError(
                f"Unable to open PDF (corrupted or invalid format): {source.name}"
            ) from exc
        # Note: a .pdf-named file containing DOCX (zip) bytes does not
        # always raise here -- fitz may open it as a zero-content document.
        # That case is still caught below by the empty-text check in
        # extract(), which raises TextExtractionError. Either way the
        # mismatch is never silently misparsed into wrong candidate data.

        if doc.is_encrypted:
            doc.close()
            logger.warning("resume_pdf_encrypted", extra={"path": str(source)})
            raise CorruptedFileError(f"PDF is password-protected: {source.name}")

        blocks: list[TextBlock] = []
        text_lines: list[str] = []
        order = 0
        try:
            for page_index, page in enumerate(doc):
                # Single API call per page; plain text is derived from the
                # same blocks pass rather than calling get_text("text") too.
                raw_blocks = page.get_text("blocks")
                for raw in sorted(raw_blocks, key=lambda b: (b[1], b[0])):
                    block_text = str(raw[4]).strip()
                    if not block_text:
                        continue
                    blocks.append(
                        TextBlock(
                            text=block_text,
                            source="body",
                            page_index=page_index,
                            order=order,
                            is_isolated=True,
                        )
                    )
                    text_lines.append(block_text)
                    order += 1
        finally:
            page_count = doc.page_count
            doc.close()

        return ExtractedText(
            plain_text="\n".join(text_lines),
            blocks=blocks,
            page_count=page_count,
        )

    def _extract_docx(self, source: Path) -> ExtractedText:
        """Extract text from a DOCX: paragraphs, tables, headers, footers.

        Args:
            source: Path to the DOCX file.

        Returns:
            Extracted text with paragraph-level block metadata.

        Raises:
            CorruptedFileError: If the file is not a valid DOCX package
                (corrupted, or an extension/content mismatch such as a
                ``.docx`` file that is actually PDF bytes).
        """
        try:
            document = Document(str(source))
        except (PackageNotFoundError, KeyError, ValueError) as exc:
            logger.warning("resume_docx_open_failed", extra={"path": str(source)})
            raise CorruptedFileError(
                f"Unable to open DOCX (corrupted or invalid format): {source.name}"
            ) from exc

        blocks: list[TextBlock] = []
        order = 0

        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                blocks.append(
                    TextBlock(text=text, source="body", order=order, is_isolated=True)
                )
                order += 1

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        blocks.append(
                            TextBlock(
                                text=text,
                                source="table",
                                order=order,
                                is_isolated=False,
                            )
                        )
                        order += 1

        for section in document.sections:
            for region, label in (
                (section.header, "header"),
                (section.footer, "footer"),
            ):
                for para in region.paragraphs:
                    text = para.text.strip()
                    if text:
                        blocks.append(
                            TextBlock(
                                text=text, source=label, order=order, is_isolated=True
                            )
                        )
                        order += 1

        plain_text = "\n".join(b.text for b in blocks)
        return ExtractedText(plain_text=plain_text, blocks=blocks, page_count=1)
